from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from citation_core import classify_and_parse, normalize_typography
from docx_io import DocumentRejected, extract_entries, preflight, sha256, write_output


def make_fixture(path: Path) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("测试论文", 0)
    p = doc.add_paragraph("正文中的论述")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("张新宝：《侵权责任法》 （第5版），中国人民大学出版社 2020 年版。")
    base = path.with_suffix(".base.docx")
    doc.save(base)
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    PR = "http://schemas.openxmlformats.org/package/2006/relationships"
    ET.register_namespace("w", W)
    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = ET.fromstring(data)
                paragraphs = root.findall(f".//{{{W}}}body/{{{W}}}p")
                run = ET.SubElement(paragraphs[1], f"{{{W}}}r")
                ET.SubElement(run, f"{{{W}}}footnoteReference", {f"{{{W}}}id": "1"})
                # The reference paragraph uses Word automatic numbering; the
                # number is not part of its text and must survive rewriting.
                ppr = paragraphs[3].find(f"{{{W}}}pPr")
                if ppr is None:
                    ppr = ET.Element(f"{{{W}}}pPr")
                    paragraphs[3].insert(0, ppr)
                numpr = ET.SubElement(ppr, f"{{{W}}}numPr")
                ET.SubElement(numpr, f"{{{W}}}ilvl", {f"{{{W}}}val": "0"})
                ET.SubElement(numpr, f"{{{W}}}numId", {f"{{{W}}}val": "5"})
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename == "[Content_Types].xml":
                addition = b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                data = data.replace(b"</Types>", addition + b"</Types>")
            elif item.filename == "word/_rels/document.xml.rels":
                addition = b'<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                data = data.replace(b"</Relationships>", addition + b"</Relationships>")
            zout.writestr(item, data)
        footnotes = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{W}">
  <w:footnote w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
  <w:footnote w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> 张新宝：《侵权责任法》 （第5版），中国人民大学出版社 2020 年版，第 73 — 75 页。</w:t></w:r></w:p></w:footnote>
</w:footnotes>'''.encode()
        zout.writestr("word/footnotes.xml", footnotes)
    base.unlink()


class CitationCoreTests(unittest.TestCase):
    def test_normalization_is_conservative(self):
        value, changes = normalize_typography("张新宝：《侵权责任法》  ，第 73 — 75 页。")
        self.assertIn("第73-75页", value)
        self.assertTrue(changes)

    def test_parse_chinese_book(self):
        entry = classify_and_parse("footnote:1", "脚注", "1", "张新宝：《侵权责任法》，中国人民大学出版社2020年版，第73-75页。")
        self.assertEqual(entry.citation_type, "zh_book_or_chapter")
        self.assertEqual(entry.title, "侵权责任法")
        self.assertEqual(entry.year, "2020")

    def test_external_styles_convert_to_legal_format(self):
        book = classify_and_parse("x", "脚注", "1", "[1] 张新宝. 侵权责任法[M]. 北京: 中国人民大学出版社, 2021:73-75.")
        self.assertIn("《侵权责任法》", book.normalized)
        self.assertIn("第73-75页", book.normalized)
        law = classify_and_parse("x", "脚注", "2", "《中华人民共和国民法典》[Z].2020-05-28.")
        self.assertIn("公布", law.normalized)
        article = classify_and_parse("x", "脚注", "3", "赵虚构. (2029). 算法平台责任的月球法解释. 火星法学, 12(4), 1-9.")
        self.assertIn("载《火星法学》", article.normalized)

    def test_incomplete_citation_is_preserved_and_reports_missing_fields(self):
        original = "赵某：《数据法研究》，载《法学》2024年。"
        entry = classify_and_parse("footnote:1", "脚注", "1", original)
        self.assertEqual(entry.citation_type, "zh_journal_article")
        self.assertEqual(entry.normalized, original)
        self.assertFalse(entry.auto_changes)
        self.assertEqual(entry.format_status, "未格式化（信息不完整）")
        self.assertEqual(entry.missing_fields, ["期号"])
        self.assertIn("期号", entry.review_note)


class DocxTests(unittest.TestCase):
    def test_extract_and_roundtrip(self):
        with tempfile.TemporaryDirectory() as td:
            src, out = Path(td) / "paper.docx", Path(td) / "out.docx"
            make_fixture(src)
            original_hash = sha256(src)
            preflight(src)
            located, trees = extract_entries(src)
            self.assertEqual({x.entry.location for x in located}, {"脚注", "参考文献"})
            write_output(src, out, located, trees)
            self.assertTrue(out.exists())
            self.assertEqual(original_hash, sha256(src))
            with zipfile.ZipFile(out) as z:
                output_document = ET.fromstring(z.read("word/document.xml"))
                output_paragraphs = output_document.findall(f".//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}body/{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}p")
                self.assertEqual(
                    ET.tostring(output_paragraphs[3].find(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}pPr/{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}numPr")),
                    b'<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:ilvl w:val="0" /><w:numId w:val="5" /></w:numPr>',
                )
                footnotes = ET.fromstring(z.read("word/footnotes.xml"))
                text_nodes = footnotes.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                self.assertTrue(any((node.text or "").startswith(" ") for node in text_nodes))

    def test_reference_automatic_numbering_is_preserved(self):
        with tempfile.TemporaryDirectory() as td:
            src, out = Path(td) / "numbered.docx", Path(td) / "numbered-out.docx"
            make_fixture(src)
            located, trees = extract_entries(src)
            reference = next(x for x in located if x.entry.location == "参考文献")
            self.assertFalse(reference.entry.original.startswith("1"))
            self.assertNotEqual(reference.entry.original, reference.entry.normalized)
            write_output(src, out, located, trees)

            def numbering(path: Path) -> tuple[str, str]:
                with zipfile.ZipFile(path) as z:
                    root = ET.fromstring(z.read("word/document.xml"))
                paragraphs = root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
                numpr = paragraphs[3].find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
                ilvl = numpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                numid = numpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                key = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                return ilvl.attrib[key], numid.attrib[key]

            self.assertEqual(numbering(src), ("0", "5"))
            self.assertEqual(numbering(out), numbering(src))

    def test_rejects_tracked_changes(self):
        with tempfile.TemporaryDirectory() as td:
            src, changed = Path(td) / "paper.docx", Path(td) / "tracked.docx"
            make_fixture(src)
            with zipfile.ZipFile(src) as zin, zipfile.ZipFile(changed, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        data = data.replace(b"</w:body>", b'<w:ins w:id="1"><w:r><w:t>x</w:t></w:r></w:ins></w:body>')
                    zout.writestr(item, data)
            with self.assertRaises(DocumentRejected):
                preflight(changed)

    def test_rejects_endnotes(self):
        with tempfile.TemporaryDirectory() as td:
            src, changed = Path(td) / "paper.docx", Path(td) / "endnotes.docx"
            make_fixture(src)
            with zipfile.ZipFile(src) as zin, zipfile.ZipFile(changed, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    zout.writestr(item, zin.read(item.filename))
                W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                zout.writestr("word/endnotes.xml", f'<w:endnotes xmlns:w="{W}"><w:endnote w:id="1"><w:p><w:r><w:t>citation</w:t></w:r></w:p></w:endnote></w:endnotes>')
            with self.assertRaises(DocumentRejected):
                preflight(changed)


if __name__ == "__main__":
    unittest.main()
