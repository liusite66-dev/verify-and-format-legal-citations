from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ET.register_namespace("w", W)


def add_footnote_reference(paragraph, note_id: int) -> None:
    run = paragraph.add_run()
    run._r.set(qn("w:rsidR"), "00000001")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "FootnoteReference")
    rpr.append(rstyle)
    run._r.append(rpr)
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), str(note_id))
    run._r.append(reference)


def build(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("算法治理中的权利保护与责任分配")
    run.bold, run.font.size = True, Pt(18)
    run.font.name = "Heiti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(31, 77, 120)

    author = doc.add_paragraph("测试作者（本文为虚构演示样例）")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(20)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Pt(22)
    body.add_run("生成式人工智能进入司法与企业治理后，算法透明、个人信息保护与侵权责任之间出现新的制度张力。")
    add_footnote_reference(body, 1)
    body.add_run("现行法律虽已确立人格权益和数据处理边界，但自动化决策的可解释性、平台注意义务及损害证明仍有争议。")
    add_footnote_reference(body, 2)
    body.add_run("未来应在技术审计、程序参与和责任分配之间形成协调机制。")
    add_footnote_reference(body, 3)
    add_footnote_reference(body, 4)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)
    hr = heading.add_run("参考文献")
    hr.bold, hr.font.size = True, Pt(16)
    hr.font.color.rgb = RGBColor(46, 116, 181)

    references = [
        "[1] 张新宝. 侵权责任法[M]. 北京: 中国人民大学出版社, 2021:73-75.",
        "赵虚构. (2029). 算法平台责任的月球法解释. 火星法学, 12(4), 1-9.",
        "《中华人民共和国民法典》[Z].2020-05-28.",
        "Jumper, J., et al. (2020). Highly accurate protein structure prediction with AlphaFold. Nature. https://doi.org/10.1038/s41586-021-03819-2.",
    ]
    for text in references:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(6)

    base = output.with_suffix(".base.docx")
    doc.save(base)

    footnotes = [
        "[1] 张新宝. 侵权责任法[M]. 北京: 中国人民大学出版社, 2021:73-75.",
        "赵虚构. (2029). 算法平台责任的月球法解释. 火星法学, 12(4), 1-9.",
        "《中华人民共和国民法典》[Z].2020-05-28.",
        "Jumper, J., et al. (2020). Highly accurate protein structure prediction with AlphaFold. Nature. https://doi.org/10.1038/s41586-021-03819-2.",
    ]
    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                addition = b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                data = data.replace(b"</Types>", addition + b"</Types>")
            elif item.filename == "word/_rels/document.xml.rels":
                addition = b'<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                data = data.replace(b"</Relationships>", addition + b"</Relationships>")
            zout.writestr(item, data)
        root = ET.Element(f"{{{W}}}footnotes")
        for note_id, note_text in [(-1, ""), (0, "")]:
            note = ET.SubElement(root, f"{{{W}}}footnote", {f"{{{W}}}id": str(note_id)})
            p = ET.SubElement(note, f"{{{W}}}p")
            r = ET.SubElement(p, f"{{{W}}}r")
            ET.SubElement(r, f"{{{W}}}{'separator' if note_id == -1 else 'continuationSeparator'}")
        for note_id, note_text in enumerate(footnotes, 1):
            note = ET.SubElement(root, f"{{{W}}}footnote", {f"{{{W}}}id": str(note_id)})
            p = ET.SubElement(note, f"{{{W}}}p")
            r1 = ET.SubElement(p, f"{{{W}}}r")
            rpr = ET.SubElement(r1, f"{{{W}}}rPr")
            ET.SubElement(rpr, f"{{{W}}}rStyle", {f"{{{W}}}val": "FootnoteReference"})
            ET.SubElement(r1, f"{{{W}}}footnoteRef")
            r2 = ET.SubElement(p, f"{{{W}}}r")
            t = ET.SubElement(r2, f"{{{W}}}t", {"{http://www.w3.org/XML/1998/namespace}space": "preserve"})
            t.text = note_text
        zout.writestr("word/footnotes.xml", ET.tostring(root, encoding="utf-8", xml_declaration=True))
    base.unlink()


if __name__ == "__main__":
    import sys
    build(Path(sys.argv[1]))
